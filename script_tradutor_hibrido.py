# -*- coding: utf-8 -*-
# === Script Tradutor Híbrido ===
# Usa googletrans para tradução rápida (substituição no texto)
# Usa Gemini API para gerar frases de exemplo contextuais
# *** NOVO: Envia e-mail de notificação ao final ***

# Bibliotecas Padrão
import argparse
import os
import re
import time
from collections import Counter
from dotenv import load_dotenv
import traceback # Para log de erros detalhado

# Bibliotecas de E-mail (NOVO)
import smtplib
import ssl
from email.message import EmailMessage

# Bibliotecas de Terceiros
import docx
from googletrans import Translator
import google.generativeai as genai
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from tqdm import tqdm

# --- Configuração do Logging (Simplificado para este script) ---
# (O script principal já faz logging detalhado, aqui focamos em printar)
print("--- Iniciando Script Tradutor Híbrido ---")

# --- Configurar NLTK ---
try:
    nltk.data.find('tokenizers/punkt')
except nltk.downloader.DownloadError:
    print("Baixando recurso 'punkt' do NLTK...")
    nltk.download('punkt', quiet=False)
try:
    nltk.data.find('corpora/stopwords')
except nltk.downloader.DownloadError:
    print("Baixando recurso 'stopwords' do NLTK...")
    nltk.download('stopwords', quiet=False)

# --- Carregar Variáveis de Ambiente ---
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
# Variáveis de E-mail (NOVO)
EMAIL_SENDER_ADDRESS = os.getenv("EMAIL_SENDER_ADDRESS")
EMAIL_SENDER_APP_PASSWORD = os.getenv("EMAIL_SENDER_APP_PASSWORD")
EMAIL_RECIPIENT_ADDRESS = os.getenv("EMAIL_RECIPIENT_ADDRESS") # Ou defina um padrão
EMAIL_SMTP_SERVER = os.getenv("EMAIL_SMTP_SERVER", "smtp.gmail.com")
EMAIL_SMTP_PORT = int(os.getenv("EMAIL_SMTP_PORT", 587))

email_configured = bool(EMAIL_SENDER_ADDRESS and EMAIL_SENDER_APP_PASSWORD and EMAIL_RECIPIENT_ADDRESS)
if not email_configured:
    print("AVISO: Configurações de e-mail (EMAIL_SENDER_ADDRESS, EMAIL_SENDER_APP_PASSWORD, EMAIL_RECIPIENT_ADDRESS) incompletas no .env. Notificação por e-mail desativada para este script.")

# --- Configuração Gemini ---
gemini_model = None
if not GOOGLE_API_KEY:
    print("AVISO: GOOGLE_API_KEY não encontrada no .env. Geração de exemplos com Gemini falhará.")
else:
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
        safety_settings_lenient = {
            'HATE': 'BLOCK_NONE', 'HARASSMENT': 'BLOCK_NONE',
            'SEXUAL' : 'BLOCK_NONE', 'DANGEROUS' : 'BLOCK_NONE'
        }
        generation_config_examples = genai.GenerationConfig(temperature=0.7)
        gemini_model = genai.GenerativeModel(
            "gemini-1.5-pro",
            safety_settings=safety_settings_lenient,
            generation_config=generation_config_examples
        )
        print(f"Modelo Gemini '{gemini_model.model_name}' inicializado para gerar exemplos.")
    except Exception as e:
        print(f"ERRO ao inicializar modelo Gemini para exemplos: {e}. Geração de exemplos desativada.")
        gemini_model = None

# --- Função de Chamada Gemini (para Exemplos) ---
def call_gemini_api_for_example(prompt, retry_count=3, base_wait=2):
    """Chama a API Gemini especificamente para gerar uma frase de exemplo."""
    if not gemini_model:
        print("  -> Modelo Gemini não disponível. Pulando geração de exemplo.")
        return None
    # ... (lógica interna da função inalterada da versão anterior) ...
    for attempt in range(retry_count):
        try:
            response = gemini_model.generate_content(prompt)
            if response.candidates and hasattr(response.candidates[0], 'content') and hasattr(response.candidates[0].content, 'parts'):
                 result_text = "".join(part.text for part in response.candidates[0].content.parts if hasattr(part, 'text')).strip()
                 if result_text:
                     result_text = re.sub(r'^["\']|["\']$', '', result_text)
                     result_text = result_text.replace("*", "")
                     return result_text.strip()
            elif hasattr(response, 'text') and response.text:
                  return response.text.strip()
            print(f"  -> Tentativa {attempt+1}: Resposta API para exemplo vazia ou inesperada.")
        except Exception as e:
            print(f"  -> Tentativa {attempt+1}: Erro ao chamar API Gemini para exemplo: {e}")
            if "RESOURCE_EXHAUSTED" in str(e) or "429" in str(e):
                 wait_time = base_wait * (2 ** attempt)
                 print(f"  -> Erro de cota, esperando {wait_time}s...")
                 time.sleep(wait_time)
            elif attempt < retry_count - 1:
                 wait_time = base_wait * (attempt + 1)
                 print(f"  -> Esperando {wait_time}s antes de tentar novamente...")
                 time.sleep(wait_time)
    print(f"  -> Falha ao gerar exemplo via API após {retry_count} tentativas.")
    return None

# --- Funções Auxiliares (DOCX, Palavras, Tradução Google, Substituição) ---

def read_docx(file_path):
    """Lê o arquivo .docx e retorna o texto completo."""
    try:
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        print(f"Texto extraído de '{os.path.basename(file_path)}' ({len(full_text)} parágrafos).")
        return ' '.join(full_text)
    except Exception as e:
        print(f"ERRO CRÍTICO ao ler o arquivo DOCX de entrada '{file_path}': {e}")
        return None

def is_valid_word(word):
    """Verifica se a palavra é válida (alfanumérica PT, > 1 char)."""
    return bool(re.match(r'^[A-Za-zÀ-ÿ0-9]+$', word)) and len(word) > 1

def get_top_words(text, n=100):
    """Obtém as n palavras mais comuns do texto, excluindo stopwords."""
    print(f"Analisando as {n} palavras mais comuns (PT)...")
    if not text:
        print("Texto de entrada vazio, não é possível encontrar palavras.")
        return []
    tokens = word_tokenize(text.lower(), language='portuguese')
    try:
        stop_words = set(stopwords.words('portuguese'))
    except Exception as e:
        print(f"ERRO ao carregar stopwords em português: {e}. Tentando continuar sem stopwords.")
        stop_words = set()
    tokens_filtered = [word for word in tokens if is_valid_word(word) and word not in stop_words]
    if not tokens_filtered:
        print("Nenhuma palavra válida (não-stopword) encontrada no texto.")
        return []
    word_counts = Counter(tokens_filtered)
    top_n = word_counts.most_common(n)
    print(f"Encontradas {len(tokens_filtered)} palavras válidas. Top {len(top_n)} selecionadas.")
    return top_n

def translate_word_with_google(word_pt):
    """Traduz uma única palavra PT -> EN usando googletrans."""
    try:
        translator = Translator()
        translation = translator.translate(word_pt, src='pt', dest='en')
        if translation and translation.text:
            if translation.text.strip().lower() != word_pt.lower():
                return translation.text.strip()
            else: return None
        else: return None
    except Exception as e:
        print(f"    -> Erro ao traduzir '{word_pt}' com googletrans: {e}")
        return None

def replace_word_in_doc(doc, word_pt, word_en):
    """Substitui a palavra PT pela EN (em negrito) no documento DOCX."""
    replacements_in_doc = 0
    pattern = rf'\b({re.escape(word_pt)})\b'
    for para in doc.paragraphs:
        if not re.search(pattern, para.text, re.IGNORECASE): continue
        runs_original = list(para.runs)
        para.clear()
        current_run_format = {'bold': False, 'italic': False, 'underline': False}
        if runs_original:
            current_run_format = {'bold': runs_original[0].bold, 'italic': runs_original[0].italic, 'underline': runs_original[0].underline}
        full_para_text_original = "".join(run.text for run in runs_original)
        parts = re.split(pattern, full_para_text_original, flags=re.IGNORECASE)
        idx_run_original = 0
        char_count_processed = 0
        for part in parts:
            if not part: continue
            part_lower = part.lower()
            is_target_word = (part_lower == word_pt.lower())
            run_format_to_apply = current_run_format
            try:
                temp_char_count = char_count_processed
                for run_idx, run_orig in enumerate(runs_original[idx_run_original:], start=idx_run_original):
                    if temp_char_count + len(run_orig.text) > char_count_processed:
                         run_format_to_apply = {'bold': run_orig.bold, 'italic': run_orig.italic, 'underline': run_orig.underline}
                         idx_run_original = run_idx
                         break
                    temp_char_count += len(run_orig.text)
            except Exception: pass
            text_to_add = word_en if is_target_word else part
            new_run = para.add_run(text_to_add)
            if is_target_word:
                new_run.bold = True
                new_run.italic = run_format_to_apply.get('italic', False)
                new_run.underline = run_format_to_apply.get('underline', False)
                replacements_in_doc += 1
            else:
                new_run.bold = run_format_to_apply.get('bold', False)
                new_run.italic = run_format_to_apply.get('italic', False)
                new_run.underline = run_format_to_apply.get('underline', False)
            if not is_target_word: current_run_format = run_format_to_apply
            char_count_processed += len(part)
    return doc

def generate_example_sentence_with_gemini(pt_word, en_word):
    """Gera uma frase de exemplo para o par PT/EN usando Gemini."""
    if not gemini_model: return None
    prompt = f"Generate one simple English example sentence using the word '{en_word}' (which is the translation of the Portuguese word '{pt_word}'). Make the sentence easy to understand for a Portuguese speaker learning English. Return ONLY the sentence itself, without any quotes or labels."
    sentence = call_gemini_api_for_example(prompt)
    if sentence:
        sentence = ' '.join(sentence.split())
        if not re.search(r'[.!?]$', sentence): sentence += '.'
        return sentence
    else: return None

# --- Função de Envio de E-mail (NOVO) ---
def send_translation_email(status_ok, input_file, output_file, words_targeted, words_translated_count, examples_generated_count, duration_secs):
    """Envia um e-mail notificando a conclusão da tradução de UM livro."""
    if not email_configured:
        print("Envio de e-mail desativado (configuração incompleta).")
        return

    status_str = "SUCESSO" if status_ok else "FALHA"
    subject = f"Tradução Híbrida Concluída ({status_str}): {os.path.basename(input_file)}"

    body = f"""
Olá,

O script de tradução híbrida concluiu o processamento do arquivo:
{os.path.basename(input_file)}

Status: {status_str}

Detalhes:
--------------------------------------------------
- Arquivo de Entrada: {input_file}
- Arquivo de Saída: {output_file if status_ok else 'N/A (Falha)'}
- Palavras Alvo para Tradução: {words_targeted}
- Palavras Efetivamente Traduzidas (Google): {words_translated_count}
- Frases de Exemplo Geradas (Gemini): {examples_generated_count}
- Duração da Tradução: {duration_secs:.2f} segundos
--------------------------------------------------

Atenciosamente,
Script Tradutor Híbrido
"""

    message = EmailMessage()
    message['Subject'] = subject
    message['From'] = EMAIL_SENDER_ADDRESS
    message['To'] = EMAIL_RECIPIENT_ADDRESS # Pode ser uma lista também: [addr1, addr2]
    message.set_content(body)

    context = ssl.create_default_context()
    server = None # Inicializa para o finally

    print(f"Tentando enviar e-mail de notificação para {EMAIL_RECIPIENT_ADDRESS}...")
    try:
        if EMAIL_SMTP_PORT == 465: # SSL
            server = smtplib.SMTP_SSL(EMAIL_SMTP_SERVER, EMAIL_SMTP_PORT, context=context, timeout=30)
            server.login(EMAIL_SENDER_ADDRESS, EMAIL_SENDER_APP_PASSWORD)
        else: # TLS (porta 587 ou outra)
            server = smtplib.SMTP(EMAIL_SMTP_SERVER, EMAIL_SMTP_PORT, timeout=30)
            server.ehlo()
            server.starttls(context=context)
            server.ehlo()
            server.login(EMAIL_SENDER_ADDRESS, EMAIL_SENDER_APP_PASSWORD)

        server.send_message(message)
        print(f"✅ E-mail de notificação de tradução enviado com sucesso!")

    except smtplib.SMTPAuthenticationError:
        print("ERRO DE AUTENTICAÇÃO do e-mail. Verifique EMAIL_SENDER_ADDRESS e EMAIL_SENDER_APP_PASSWORD.")
        print("Lembre-se: Para Gmail com 2FA, use uma 'Senha de App'.")
    except Exception as e:
        print(f"ERRO ao enviar e-mail de notificação: {e}")
        traceback.print_exc()
    finally:
        if server:
            try:
                server.quit()
            except Exception: pass # Ignora erros ao fechar

# === Função Principal do Script Tradutor ===
def process_translation_hybrid(input_path, output_path, num_words=100):
    """
    Função principal: lê DOCX, traduz palavras comuns (Google),
    substitui no texto, gera exemplos (Gemini) e salva novo DOCX.
    MODIFICADO: Retorna status e contagens para o e-mail.
    Retorna: (bool_success, words_translated_count, examples_generated_count)
    """
    print(f"\n--- Iniciando Processo de Tradução Híbrida ---")
    print(f"Arquivo de Entrada: {input_path}")
    print(f"Arquivo de Saída: {output_path}")
    print(f"Número de palavras para processar: {num_words}")

    words_translated_count = 0
    examples_generated_count = 0

    # 1. Ler DOCX e Extrair Texto
    print("\n[Passo 1/6] Lendo DOCX e extraindo texto...")
    full_text = read_docx(input_path)
    if not full_text:
        print("ERRO CRÍTICO: Falha ao ler texto do arquivo de entrada. Abortando.")
        return False, 0, 0 # Retorna falha e contagens zeradas

    # 2. Identificar Palavras Comuns
    print("\n[Passo 2/6] Identificando palavras comuns...")
    top_words_pt = get_top_words(full_text, num_words)
    if not top_words_pt:
        print("Nenhuma palavra comum encontrada para traduzir. Saindo.")
        return True, 0, 0 # Terminou "com sucesso", mas sem fazer nada

    # 3. Traduzir Palavras com Googletrans
    print("\n[Passo 3/6] Traduzindo palavras com Google Translate...")
    translated_words = {}
    word_pairs_for_examples = []
    with tqdm(total=len(top_words_pt), desc="Traduzindo (Google)", unit="palavra") as pbar:
        for word, count in top_words_pt:
            translation = translate_word_with_google(word)
            if translation:
                translated_words[word] = translation
                word_pairs_for_examples.append({'pt': word, 'en': translation})
            pbar.update(1)

    words_translated_count = len(translated_words) # Conta quantas realmente foram traduzidas

    if words_translated_count == 0:
        print("Nenhuma palavra foi traduzida com sucesso pelo Google Translate. Abortando modificações.")
        return True, 0, 0 # Terminou "com sucesso", mas sem fazer nada

    print(f"\n{words_translated_count} palavras traduzidas com sucesso via Google Translate.")

    # 4. Carregar DOCX e Substituir Palavras
    print("\n[Passo 4/6] Carregando documento para substituição...")
    try:
        doc = docx.Document(input_path)
    except Exception as e:
        print(f"ERRO CRÍTICO ao carregar DOCX '{input_path}' para modificação: {e}")
        return False, words_translated_count, 0

    print("Substituindo palavras no documento (pode levar um tempo)...")
    with tqdm(total=len(translated_words), desc="Substituindo no DOCX", unit="palavra") as pbar:
        for word_pt, word_en in translated_words.items():
            doc = replace_word_in_doc(doc, word_pt, word_en)
            pbar.update(1)

    # 5. Gerar Frases de Exemplo com Gemini
    print("\n[Passo 5/6] Gerando frases de exemplo com Gemini API...")
    example_sentences_data = []
    if gemini_model:
        with tqdm(total=len(word_pairs_for_examples), desc="Gerando Exemplos (Gemini)", unit="frase") as pbar:
            for pair in word_pairs_for_examples:
                sentence = generate_example_sentence_with_gemini(pair['pt'], pair['en'])
                if sentence:
                    example_sentences_data.append({
                        'pt_word': pair['pt'],
                        'en_word': pair['en'],
                        'sentence': sentence
                    })
                pbar.update(1)
        examples_generated_count = len(example_sentences_data) # Conta quantos exemplos foram gerados
        print(f"\n{examples_generated_count} frases de exemplo geradas via Gemini.")
    else:
        print("Geração de exemplos pulada (Modelo Gemini não disponível).")
        examples_generated_count = 0

    # 6. Adicionar Listas e Salvar
    print("\n[Passo 6/6] Adicionando listas ao final do documento e salvando...")
    try:
        doc.add_paragraph("\n--- Palavras Traduzidas (PT -> EN) ---").bold = True
        for word_pt in sorted(translated_words.keys()):
            word_en = translated_words[word_pt]
            para = doc.add_paragraph()
            para.add_run(f"{word_pt} - ")
            run_en = para.add_run(word_en); run_en.bold = True

        if example_sentences_data:
             doc.add_paragraph("\n--- Frases de Exemplo (Gemini API) ---").bold = True
             for item in sorted(example_sentences_data, key=lambda x: x['pt_word']):
                 para = doc.add_paragraph()
                 run_words = para.add_run(f"{item['pt_word']} ({item['en_word']}): "); run_words.bold = True
                 para.add_run(item['sentence'])

        doc.save(output_path)
        print(f"\nArquivo traduzido salvo com sucesso em: {output_path}")
        return True, words_translated_count, examples_generated_count # Retorna sucesso e contagens

    except Exception as e:
        print(f"ERRO CRÍTICO ao adicionar listas ou salvar o arquivo final '{output_path}': {e}")
        traceback.print_exc()
        return False, words_translated_count, examples_generated_count # Retorna falha e contagens parciais

# --- Bloco de Execução Principal (quando rodado diretamente) ---
if __name__ == "__main__":
    script_start_time = time.time()
    print("Executando Script Tradutor Híbrido...")

    parser = argparse.ArgumentParser(description="Traduz palavras comuns PT->EN em DOCX (Google Translate) e gera exemplos (Gemini API).")
    parser.add_argument("--input", required=True, help="Caminho para o arquivo DOCX de entrada (corrigido).")
    parser.add_argument("--output", required=True, help="Caminho para salvar o arquivo DOCX de saída (traduzido).")
    parser.add_argument("--words", type=int, default=100, help="Número de palavras mais comuns a traduzir e gerar exemplos.")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"ERRO FATAL: Arquivo de entrada não encontrado: {args.input}")
        exit(1) # Sai com código de erro para o subprocesso

    # Chama a função principal e captura os resultados
    success, final_words_translated, final_examples_generated = process_translation_hybrid(
        args.input, args.output, args.words
    )

    script_end_time = time.time()
    script_duration = script_end_time - script_start_time

    # Envia e-mail de notificação (NOVO)
    send_translation_email(
        status_ok=success,
        input_file=args.input,
        output_file=args.output,
        words_targeted=args.words,
        words_translated_count=final_words_translated,
        examples_generated_count=final_examples_generated,
        duration_secs=script_duration
    )

    # Retorna código de saída apropriado
    if success:
        print(f"\n--- Processo de Tradução Híbrida concluído com SUCESSO em {script_duration:.2f}s. ---")
        exit(0) # Código 0 indica sucesso
    else:
        print(f"\n--- Processo de Tradução Híbrida concluído com ERROS em {script_duration:.2f}s. ---")
        exit(1) # Código 1 indica falha